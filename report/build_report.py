from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Dict, List
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = PROJECT_ROOT / "heart.csv"
OUTPUT_PATH = PROJECT_ROOT / "report" / "report_submission.docx"
SRC_PATH = PROJECT_ROOT / "src"

if str(SRC_PATH) not in sys.path:
    sys.path.insert(0, str(SRC_PATH))

from analysis import generate_plots, probability_summary, train_logistic_regression

TEAM_NAME = "CardioSense"
PROJECT_TITLE = "Heart Disease Risk Prediction and Statistical Analysis System"
LEADER = "Husnain Khalid"
SUBMISSION_DATE = date(2026, 5, 6)

ACCENT_COLOR = RGBColor(15, 118, 110)
ACCENT_HEX = "0F766E"

MEMBERS: List[Dict[str, str]] = [
    {"roll": "22f-3190", "name": "________", "section": "______"},
    {"roll": "22f-8758", "name": "________", "section": "______"},
    {"roll": "22f-3812", "name": "________", "section": "______"},
    {"roll": "22f-3143", "name": "________", "section": "______"},
    {"roll": "22f-3335", "name": "________", "section": "______"},
]

DATASET_NAME = "Heart Disease Dataset (Kaggle)"
DATASET_LINK = "https://www.kaggle.com/datasets/johnsmith88/heart-disease-dataset?select=heart.csv"
VARIABLES = [
    "age",
    "sex",
    "cp",
    "trestbps",
    "chol",
    "fbs",
    "restecg",
    "thalach",
    "exang",
    "oldpeak",
    "slope",
    "ca",
    "thal",
    "target",
]

BRIEF_SUMMARY = (
    "This project uses a real-world heart disease dataset with multiple variables to perform statistical "
    "analysis, probability modeling, and predictive regression through a web-based application for heart "
    "disease risk assessment."
)


def set_style_font(style, name: str, size: int, color: RGBColor | None = None, bold: bool | None = None) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    if color is not None:
        font.color.rgb = color
    if bold is not None:
        font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_centered_text(
    doc: Document,
    text: str,
    font_name: str,
    size: int,
    color: RGBColor | None = None,
    bold: bool = False,
    spacing_after: int = 6,
) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    para.paragraph_format.space_after = Pt(spacing_after)
    run_props = run._element.get_or_add_rPr()
    run_props.rFonts.set(qn("w:eastAsia"), font_name)


def add_section_heading(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=1)
    heading.paragraph_format.space_after = Pt(6)


def add_divider(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), ACCENT_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(82, 96, 109)
    run_props = run._element.get_or_add_rPr()
    run_props.rFonts.set(qn("w:eastAsia"), "Calibri")


def build_report() -> None:
    df = pd.read_csv(DATA_PATH)
    plots = generate_plots(df, PROJECT_ROOT / "outputs")
    prob = probability_summary(df)
    _, metrics = train_logistic_regression(df)

    chol_means = df.groupby("target")["chol"].mean().to_dict()
    thalach_means = df.groupby("target")["thalach"].mean().to_dict()
    summary = {
        "Rows": df.shape[0],
        "Columns": df.shape[1],
        "Target Rate": f"{df['target'].mean() * 100:.1f}%",
    }

    conclusion_text = (
        f"The dataset is nearly balanced with {summary['Target Rate']} target=1, so the analysis reflects both outcomes. "
        f"Probability results show a higher observed target rate for females ({prob['prob_target_given_female'] * 100:.1f}%) "
        f"than males ({prob['prob_target_given_male'] * 100:.1f}%), and about {prob['cholesterol_gt_240'] * 100:.1f}% "
        "of records have cholesterol greater than 240. Visual analysis indicates higher average max heart rate "
        f"(thalach) for target=1 ({thalach_means.get(1, 0.0):.1f}) compared with target=0 "
        f"({thalach_means.get(0, 0.0):.1f}), while average cholesterol is slightly higher for target=0 "
        f"({chol_means.get(0, 0.0):.1f}) than target=1 ({chol_means.get(1, 0.0):.1f}). "
        f"The logistic regression model performs strongly (accuracy {metrics.accuracy:.3f}, precision {metrics.precision:.3f}, "
        f"recall {metrics.recall:.3f}, AUC {metrics.roc_auc:.3f}), showing the variables provide useful predictive signal. "
        "Overall, the system summarizes key risk patterns and delivers a practical early screening aid, but results "
        "should be interpreted as supportive analytics rather than a clinical diagnosis."
    )

    doc = Document()

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11)
    set_style_font(styles["Heading 1"], "Garamond", 16, ACCENT_COLOR, True)
    set_style_font(styles["Heading 2"], "Garamond", 13, ACCENT_COLOR, True)

    add_centered_text(doc, "Semester Project Report", "Garamond", 22, ACCENT_COLOR, True)
    add_centered_text(doc, PROJECT_TITLE, "Garamond", 16, bold=True)
    add_centered_text(doc, f"Team: {TEAM_NAME}", "Calibri", 12)
    add_centered_text(doc, f"Group Leader: {LEADER}", "Calibri", 12)
    add_centered_text(doc, f"Submission Date: {SUBMISSION_DATE.strftime('%d %B %Y')}", "Calibri", 11)
    add_centered_text(doc, "[Team Logo Placeholder]", "Calibri", 11)
    add_divider(doc)

    summary_table = doc.add_table(rows=1, cols=3)
    summary_table.style = "Light Shading Accent 1"
    header_cells = summary_table.rows[0].cells
    header_cells[0].text = "Records"
    header_cells[1].text = "Variables"
    header_cells[2].text = "Target Rate"
    row_cells = summary_table.add_row().cells
    row_cells[0].text = str(summary["Rows"])
    row_cells[1].text = str(summary["Columns"])
    row_cells[2].text = summary["Target Rate"]

    doc.add_page_break()

    add_section_heading(doc, "Group Members")
    member_table = doc.add_table(rows=1, cols=3)
    member_table.style = "Light Shading Accent 1"
    header_cells = member_table.rows[0].cells
    header_cells[0].text = "Roll Number"
    header_cells[1].text = "Name"
    header_cells[2].text = "Section"
    for member in MEMBERS:
        row_cells = member_table.add_row().cells
        row_cells[0].text = member["roll"]
        row_cells[1].text = member["name"]
        row_cells[2].text = member["section"]

    add_section_heading(doc, "1) Problem Statement")
    doc.add_paragraph(
        "Heart disease is a major health risk. This project analyzes a real-world dataset to identify "
        "key factors and provide a risk prediction tool using statistical methods and regression modeling."
    )

    add_section_heading(doc, "2) Objective")
    doc.add_paragraph(
        "Build a web-based system that performs statistical analysis, probability modeling, and "
        "regression-based prediction for heart disease risk, and present results in a clear visual format."
    )

    add_section_heading(doc, "3) Data Description")
    doc.add_paragraph(f"Dataset Name: {DATASET_NAME}")
    doc.add_paragraph(f"Link: {DATASET_LINK}")
    doc.add_paragraph("Variables:")
    for var in VARIABLES:
        doc.add_paragraph(var, style="List Bullet")

    add_section_heading(doc, "Brief Summary")
    doc.add_paragraph(BRIEF_SUMMARY)

    doc.add_paragraph("Dataset Summary:")
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Light Shading Accent 1"
    summary_table.rows[0].cells[0].text = "Metric"
    summary_table.rows[0].cells[1].text = "Value"
    for key, value in summary.items():
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)

    add_section_heading(doc, "4) Results")
    doc.add_paragraph(
        "Include screenshots of the application and key results, such as descriptive statistics, "
        "confidence intervals, probability summaries, and regression model metrics."
    )
    doc.add_paragraph("[Insert screenshot: Overview dashboard]")
    doc.add_paragraph("[Insert screenshot: Descriptive statistics and confidence intervals]")
    doc.add_paragraph("[Insert screenshot: Probability summary]")
    doc.add_paragraph("[Insert screenshot: Model metrics and prediction form]")

    add_section_heading(doc, "Visual Results")
    doc.add_paragraph("Charts generated from the dataset are included below.")
    doc.add_picture(str(plots["target_count"]), width=Inches(5.8))
    add_caption(doc, "Figure 1: Target count distribution")
    doc.add_picture(str(plots["correlation_heatmap"]), width=Inches(5.8))
    add_caption(doc, "Figure 2: Correlation heatmap")
    doc.add_picture(str(plots["hist_chol"]), width=Inches(5.8))
    add_caption(doc, "Figure 3: Cholesterol distribution")
    doc.add_picture(str(plots["chol_by_target"]), width=Inches(5.8))
    add_caption(doc, "Figure 4: Cholesterol by target")

    add_section_heading(doc, "5) Codes")
    code_note = doc.add_paragraph()
    code_note.add_run(
        "Paste the full source code here. Use font size 9 with line spacing 1, and add proper titles and comments."
    ).font.size = Pt(9)
    code_note.paragraph_format.line_spacing = 1.0

    add_section_heading(doc, "6) Conclusion")
    doc.add_paragraph(conclusion_text)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
